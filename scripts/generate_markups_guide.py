#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Generate a comprehensive Word document covering multiple markup languages:
- Markdown
- reStructuredText
- AsciiDoc
- HTML/XML
- LaTeX
- Mermaid / PlantUML
- YAML / JSON / TOML
- Pandoc conversion usage

Output: /workspace/Markups_Comprehensive_Guide.docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import argparse
import os
from pathlib import Path
from typing import Iterable


DEFAULT_OUTPUT_PATH = "/workspace/Markups_Comprehensive_Guide.docx"


def ensure_code_paragraph_style(document: Document, style_name: str = "Code Block") -> None:
    """Ensure a paragraph style for code blocks exists and is monospaced."""
    styles = document.styles
    if style_name in [s.name for s in styles]:
        return
    style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Courier New"
    font.size = Pt(10)


def add_toc(document: Document, title: str = "Table of Contents") -> None:
    """Insert a Table of Contents field that Word can update on open."""
    document.add_page_break()
    heading = document.add_paragraph(title)
    heading.style = document.styles["Heading 1"]
    p = document.add_paragraph()

    # Field: TOC \o "1-3" \h \z \u
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "TOC \\o \"1-3\" \\h \\z \\u"
    run._r.append(instr_text)

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    # Placeholder run for the TOC content
    run2 = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_end)

    note = document.add_paragraph("Right-click the TOC and choose 'Update Field' in Word to populate.")
    note.style = document.styles["Intense Quote"]


def add_cover(document: Document) -> None:
    """Add a simple cover page with title, subtitle, and date."""
    title = document.add_paragraph("Markups: One Comprehensive Guide")
    title_format = title.paragraph_format
    title_format.space_after = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style = document.styles["Title"]

    subtitle = document.add_paragraph("Markdown, reStructuredText, AsciiDoc, HTML/XML, LaTeX, Diagrams, Data Formats, and Pandoc")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = document.styles["Subtitle"]

    date_p = document.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()


def add_code_block(document: Document, code: str, style_name: str = "Code Block") -> None:
    """Add a multi-line code block using a monospaced paragraph style."""
    for line in code.rstrip("\n").split("\n"):
        p = document.add_paragraph(line)
        p.style = document.styles[style_name]


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_section_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph(text)
    p.style = document.styles[f"Heading {level}"]


def add_markdown_section(document: Document) -> None:
    add_section_heading(document, "Markdown", 1)
    add_bullets(document, [
        "Lightweight plain-text syntax for formatting documents.",
        "Common variants: CommonMark, GitHub Flavored Markdown (GFM).",
        "Great for READMEs, docs, notes; easy to write and diff.",
    ])
    add_section_heading(document, "Syntax Overview", 2)
    add_code_block(document, """
# Heading 1
## Heading 2

**bold** _italic_ `inline code`

- List item
1. Numbered item

```python
def hello(name):
    print(f"Hello {name}")
```

> Quote

| Col A | Col B |
|-------|-------|
|  1    |  2    |
""".strip("\n"))
    add_section_heading(document, "Pros & Cons", 2)
    add_bullets(document, [
        "Pros: ubiquitous, simple, great tooling.",
        "Cons: dialect differences, limited advanced layout without extensions.",
    ])
    add_section_heading(document, "Tools", 2)
    add_bullets(document, [
        "Renderers: Pandoc, mdBook, MkDocs, Sphinx (via MyST).",
        "Linters: markdownlint, remark-lint.",
    ])
    add_section_heading(document, "Convert to DOCX", 2)
    add_code_block(document, """
pandoc input.md -o output.docx --from=gfm --reference-doc=template.docx
""".strip("\n"))


def add_rst_section(document: Document) -> None:
    add_section_heading(document, "reStructuredText (reST)", 1)
    add_bullets(document, [
        "Powerful, extensible markup used widely in Python docs.",
        "Strong directive system and roles; integrates with Sphinx.",
    ])
    add_section_heading(document, "Syntax Overview", 2)
    add_code_block(document, """
Heading 1
=========

Heading 2
---------

**bold** *italic* ``inline code``

- Bullet list

.. code-block:: python

   def hello(name):
       print(f"Hello {name}")

.. note:: Notes and admonitions
""".strip("\n"))
    add_section_heading(document, "Pros & Cons", 2)
    add_bullets(document, [
        "Pros: feature-rich, precise, excellent for large docs.",
        "Cons: heavier syntax, steeper learning curve.",
    ])
    add_section_heading(document, "Convert to DOCX", 2)
    add_code_block(document, """
pandoc input.rst -o output.docx --from=rst
""".strip("\n"))


def add_asciidoc_section(document: Document) -> None:
    add_section_heading(document, "AsciiDoc", 1)
    add_bullets(document, [
        "Rich, semantic markup for technical documentation.",
        "Popular in enterprises; tools: Asciidoctor, Antora.",
    ])
    add_section_heading(document, "Syntax Overview", 2)
    add_code_block(document, """
= Document Title
:toc:

== Section

*bold* _italic_ `mono`

* List item
. Numbered item

[source,python]
----
def hello(name):
    print(f"Hello {name}")
----
""".strip("\n"))
    add_section_heading(document, "Pros & Cons", 2)
    add_bullets(document, [
        "Pros: expressive, modular, powerful includes and attributes.",
        "Cons: fewer hosted renderers than Markdown.",
    ])
    add_section_heading(document, "Convert to DOCX", 2)
    add_code_block(document, """
pandoc input.adoc -o output.docx --from=asciidoc
""".strip("\n"))


def add_html_xml_section(document: Document) -> None:
    add_section_heading(document, "HTML / XML", 1)
    add_bullets(document, [
        "Markup languages for structured content; HTML for web, XML generic.",
        "Precise control, but verbose to author by hand.",
    ])
    add_section_heading(document, "Examples", 2)
    add_code_block(document, """
<!-- HTML snippet -->
<article>
  <h1>Title</h1>
  <p><strong>Bold</strong> and <em>italic</em></p>
</article>

<!-- XML snippet -->
<report>
  <section title="Intro">
    <para>Hello</para>
  </section>
</report>
""".strip("\n"))
    add_section_heading(document, "Convert to DOCX", 2)
    add_code_block(document, """
pandoc input.html -o output.docx --from=html
""".strip("\n"))


def add_latex_section(document: Document) -> None:
    add_section_heading(document, "LaTeX", 1)
    add_bullets(document, [
        "Typesetting system for scientific and technical documents.",
        "Excellent math support; steep learning curve.",
    ])
    add_section_heading(document, "Example", 2)
    add_code_block(document, """
\\documentclass{article}
\\begin{document}
\\section{Intro}
Bold: \\textbf{strong}, Italic: \\emph{emphasis}.

Inline math: $E=mc^2$

Display math:
\\[
  \\int_0^1 x^2 \\; dx = \\frac{1}{3}
\\]
\\end{document}
""".strip("\n"))
    add_section_heading(document, "Convert to DOCX", 2)
    add_code_block(document, """
pandoc input.tex -o output.docx --from=latex
""".strip("\n"))


def add_diagrams_section(document: Document) -> None:
    add_section_heading(document, "Diagrams: Mermaid / PlantUML", 1)
    add_bullets(document, [
        "Text-based diagram definitions for sequence, flowchart, class, etc.",
        "Usually rendered to SVG/PNG; can be embedded via toolchains.",
    ])
    add_section_heading(document, "Mermaid", 2)
    add_code_block(document, """
```mermaid
flowchart TD
  A[Start] --> B{Choice}
  B -->|Yes| C[Do thing]
  B -->|No| D[Stop]
```
""".strip("\n"))
    add_section_heading(document, "PlantUML", 2)
    add_code_block(document, """
@startuml
Alice -> Bob: Hello
Bob --> Alice: Hi!
@enduml
""".strip("\n"))
    add_section_heading(document, "Tips", 2)
    add_bullets(document, [
        "Render diagrams to images first, then include in documents.",
        "Some toolchains (Pandoc filters) render diagrams inline.",
    ])


def add_data_formats_section(document: Document) -> None:
    add_section_heading(document, "Data Serialization: YAML / JSON / TOML", 1)
    add_bullets(document, [
        "Configuration-oriented formats often embedded in docs or pipelines.",
        "Do not confuse with presentation markup; include as code blocks.",
    ])
    add_section_heading(document, "Examples", 2)
    add_code_block(document, """
# YAML
title: Guide
authors:
  - name: Ada
  - name: Linus

// JSON
{
  "title": "Guide",
  "authors": ["Ada", "Linus"]
}

# TOML
title = "Guide"
authors = ["Ada", "Linus"]
""".strip("\n"))


def add_pandoc_section(document: Document) -> None:
    add_section_heading(document, "Converting Between Markups with Pandoc", 1)
    add_bullets(document, [
        "Pandoc converts among many formats including DOCX, PDF (via LaTeX), HTML.",
        "Use a reference DOCX to control styles.",
    ])
    add_section_heading(document, "Common Commands", 2)
    add_code_block(document, """
# Markdown to DOCX
pandoc input.md -o output.docx --from=gfm --reference-doc=template.docx

# AsciiDoc to DOCX
pandoc input.adoc -o output.docx --from=asciidoc

# reST to DOCX
pandoc input.rst -o output.docx --from=rst

# HTML to DOCX
pandoc input.html -o output.docx --from=html

# LaTeX to DOCX (structure only; complex math may rasterize)
pandoc input.tex -o output.docx --from=latex
""".strip("\n"))
    add_section_heading(document, "Reference DOCX Tips", 2)
    add_numbered(document, [
        "Create a Word template with desired styles (Heading 1..3, Code).",
        "Export as reference DOCX and pass via --reference-doc.",
        "Ensure style names match exactly for best results.",
    ])


def build_document(output_path: Path) -> None:
    document = Document()

    # Prepare styles
    ensure_code_paragraph_style(document, "Code Block")

    # Cover & TOC
    add_cover(document)
    add_toc(document)

    # Sections
    add_section_heading(document, "Introduction", 1)
    add_bullets(document, [
        "This guide summarizes common markups, their syntax, strengths, and conversion to Word.",
        "Use it as a quick reference and a starting point for toolchain selection.",
    ])

    add_markdown_section(document)
    add_rst_section(document)
    add_asciidoc_section(document)
    add_html_xml_section(document)
    add_latex_section(document)
    add_diagrams_section(document)
    add_data_formats_section(document)
    add_pandoc_section(document)

    # Closing
    add_section_heading(document, "Further Tips", 1)
    add_bullets(document, [
        "Prefer one primary markup across a project to avoid dialect drift.",
        "Automate conversion (CI) for repeatable outputs.",
        "Validate with linters and link checkers.",
    ])

    document.add_page_break()
    add_section_heading(document, "References", 1)
    add_bullets(document, [
        "CommonMark Spec",
        "Pandoc User Guide",
        "Sphinx Docs",
        "Asciidoctor Docs",
    ])

    document.save(str(output_path))
    print(f"Generated: {output_path}")


def infer_format_from_extension(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    mapping = {
        ".md": "Markdown",
        ".markdown": "Markdown",
        ".rst": "reStructuredText",
        ".adoc": "AsciiDoc",
        ".asciidoc": "AsciiDoc",
        ".html": "HTML",
        ".htm": "HTML",
        ".xml": "XML",
        ".tex": "LaTeX",
        ".mmd": "Mermaid",
        ".mermaid": "Mermaid",
        ".puml": "PlantUML",
        ".uml": "PlantUML",
        ".yaml": "YAML",
        ".yml": "YAML",
        ".json": "JSON",
        ".toml": "TOML",
    }
    return mapping.get(ext, ext.lstrip(".").upper() or "Unknown")


def iter_markup_files(root_dir: Path) -> Iterable[Path]:
    allowed_exts = {
        ".md", ".markdown", ".rst", ".adoc", ".asciidoc", ".html", ".htm",
        ".xml", ".tex", ".mmd", ".mermaid", ".puml", ".uml", ".yaml",
        ".yml", ".json", ".toml",
    }
    for base, _dirs, files in os.walk(root_dir):
        for name in sorted(files):
            p = Path(base) / name
            if p.suffix.lower() in allowed_exts:
                yield p


def add_file_section(document: Document, root_dir: Path, file_path: Path) -> None:
    rel = file_path.relative_to(root_dir)
    title = f"{rel.as_posix()} ({infer_format_from_extension(file_path)})"
    add_section_heading(document, title, 1)

    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1")

    # Keep content as a code block to preserve original markup
    add_code_block(document, text)


def build_from_folder(input_dir: Path, output_path: Path) -> None:
    document = Document()
    ensure_code_paragraph_style(document, "Code Block")

    add_cover(document)
    add_toc(document)

    add_section_heading(document, f"Bundle Overview", 1)
    add_bullets(document, [
        f"Source folder: {input_dir}",
        "Each file is included as a separate section with its raw markup.",
        "Use Pandoc commands (see guide section) to convert markups if needed.",
    ])

    count = 0
    for f in iter_markup_files(input_dir):
        add_file_section(document, input_dir, f)
        count += 1

    if count == 0:
        add_section_heading(document, "No supported markup files found.", 2)

    document.save(str(output_path))
    print(f"Generated bundle: {output_path} (files: {count})")


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate Word docs for markups.")
    parser.add_argument("--mode", choices=["guide", "compile"], default="guide",
                        help="guide: comprehensive guide; compile: include all markup files from a folder")
    parser.add_argument("--input-dir", type=str, default="",
                        help="Folder containing markup files (required for compile mode)")
    parser.add_argument("--output", type=str, default=DEFAULT_OUTPUT_PATH,
                        help="Output .docx path")
    args = parser.parse_args()

    if args.mode == "compile":
        if not args.input_dir:
            raise SystemExit("--input-dir is required for compile mode")
        input_dir = Path(args.input_dir).resolve()
        if not input_dir.exists() or not input_dir.is_dir():
            raise SystemExit(f"Input directory not found: {input_dir}")
        build_from_folder(input_dir, Path(args.output).resolve())
    else:
        build_document(Path(args.output).resolve())


if __name__ == "__main__":
    main()

